#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
md_to_docx.py
把 Markdown 使用说明转成 Word (.docx) 文档。

用法：
    python scripts/md_to_docx.py --input 使用说明.md --output 使用说明.docx
"""

import argparse
import re
from pathlib import Path

import markdown
from bs4 import BeautifulSoup
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_formatted_text(paragraph, element):
    """把 BeautifulSoup 节点里的文本按加粗/斜体/代码格式追加到段落。"""
    for child in element.children:
        if child.name is None:
            paragraph.add_run(child.string or '')
        elif child.name == 'strong':
            run = paragraph.add_run(child.get_text())
            run.bold = True
        elif child.name == 'em':
            run = paragraph.add_run(child.get_text())
            run.italic = True
        elif child.name == 'code':
            run = paragraph.add_run(child.get_text())
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x2C, 0x0C)
        else:
            add_formatted_text(paragraph, child)


def convert_md_to_docx(input_path: str, output_path: str):
    input_file = Path(input_path)
    output_file = Path(output_path)

    if not input_file.exists():
        raise FileNotFoundError(f"输入文件不存在: {input_file}")

    md_text = input_file.read_text(encoding='utf-8')
    html = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])
    soup = BeautifulSoup(html, 'html.parser')

    doc = Document()

    # 默认中文字体
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), 'Microsoft YaHei')
    doc.styles['Normal'].font.size = Pt(11)

    for elem in soup.find_all(recursive=False):
        if elem.name == 'h1':
            p = doc.add_heading(elem.get_text(), level=0)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif elem.name == 'h2':
            doc.add_heading(elem.get_text(), level=1)
        elif elem.name == 'h3':
            doc.add_heading(elem.get_text(), level=2)
        elif elem.name == 'h4':
            doc.add_heading(elem.get_text(), level=3)
        elif elem.name == 'p':
            p = doc.add_paragraph()
            add_formatted_text(p, elem)
        elif elem.name == 'ul':
            for li in elem.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_text(p, li)
        elif elem.name == 'ol':
            for li in elem.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Number')
                add_formatted_text(p, li)
        elif elem.name == 'pre':
            code = elem.get_text()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(code)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif elem.name == 'table':
            rows = elem.find_all('tr')
            if not rows:
                continue
            max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.style = 'Light Grid Accent 1'
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells):
                    table.rows[i].cells[j].text = cell.get_text()

    doc.save(str(output_file))
    print(f"已生成 Word 文档: {output_file}")


def main():
    parser = argparse.ArgumentParser(description="Markdown 转 Word")
    parser.add_argument('--input', default='使用说明.md', help='输入 Markdown 文件')
    parser.add_argument('--output', default='使用说明.docx', help='输出 Word 文件')
    args = parser.parse_args()
    convert_md_to_docx(args.input, args.output)


if __name__ == '__main__':
    main()
